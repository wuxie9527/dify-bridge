# -*- coding: utf-8 -*-
"""
报告审核 API（支持文件上传和 URL 下载）
提供三个核心接口：
1. Excel 提取（支持 UploadFile 或 URL）
2. Word 文本提取（支持 URL）
3. 批注写回（支持 UploadFile 或 URL）
"""
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from typing import Dict, Any, Optional, List, Tuple
import os
import json
import logging
import httpx
import tempfile
import uuid
from datetime import datetime
from io import BytesIO
from docx import Document

logger = logging.getLogger(__name__)

from app.services.excel_extractor import ExcelExtractor
from app.services.excel_annotator import ExcelAnnotator
from app.services.word_annotator import WordAnnotator

router = APIRouter(prefix="/api/v1/report", tags=["报告审核"])

# 临时文件存储目录
TEMP_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "uploads")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "outputs")

os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)


def save_temp_file(file: UploadFile = None, suffix: str = "", file_url: str = None, timeout: int = 120) -> str:
    """
    保存文件到临时目录（支持 UploadFile 或 URL 下载）

    Args:
        file: UploadFile 对象（可选）
        suffix: 文件后缀
        file_url: 文件 URL（可选，与 file 二选一）
        timeout: 下载超时时间（秒），默认 120 秒（大文件可能需要更长时间）

    Returns:
        保存后的文件路径
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if file:
        # 传统文件上传
        filename = f"{timestamp}_{file.filename}{suffix}"
        file_path = os.path.join(TEMP_DIR, filename)
        with open(file_path, "wb") as f:
            f.write(file.file.read())
        return file_path

    elif file_url:
        # 从 URL 下载（增加超时）
        file_path = os.path.join(TEMP_DIR, f"{timestamp}_download{suffix}")
        with httpx.Client(timeout=timeout) as client:
            resp = client.get(file_url)
            resp.raise_for_status()
            with open(file_path, "wb") as f:
                f.write(resp.content)
        return file_path

    else:
        raise ValueError("必须提供 file 或 file_url")


def validate_annotations(audit_data: Dict, files_provided: Dict[str, bool]):
    """
    校验文件和批注的匹配关系

    Args:
        audit_data: 审核结果 JSON 对象
        files_provided: 文件提供情况 {"excel": True/False, "report": True/False, "explanation": True/False}

    Raises:
        HTTPException: 校验失败时抛出
    """
    annotations = audit_data.get("annotations", {})

    errors = []

    # Excel 校验
    if files_provided.get("excel", False):
        excel_annotations = annotations.get("excel", [])
        if not excel_annotations:
            errors.append("上传了 Excel 文件，但 annotations.excel 为空。请提供 Excel 单元格的批注内容。")
        else:
            # 验证每个 Excel 批注的 location 格式
            for i, ann in enumerate(excel_annotations):
                location = ann.get("location", "")
                if "!" not in location:
                    errors.append(f"Excel 批注[{i}] 的 location 格式错误：'{location}'。应为 'Sheet 名!单元格' 格式，如 '资产明细表!C3'")
                if not ann.get("description"):
                    errors.append(f"Excel 批注[{i}] 缺少 description 字段")
                if not ann.get("suggestion"):
                    errors.append(f"Excel 批注[{i}] 缺少 suggestion 字段")

    # Word 报告校验
    if files_provided.get("report", False):
        report_annotations = annotations.get("report", [])
        if not report_annotations:
            errors.append("上传了评估报告 Word 文件，但 annotations.report 为空。请提供报告章节的批注内容。")
        else:
            for i, ann in enumerate(report_annotations):
                original_text = ann.get("original_text", "")
                if not original_text:
                    errors.append(f"报告批注[{i}] 缺少 original_text 字段（应包含报告中的原文段落）")
                if not ann.get("description"):
                    errors.append(f"报告批注[{i}] 缺少 description 字段")
                if not ann.get("suggestion"):
                    errors.append(f"报告批注[{i}] 缺少 suggestion 字段")

    # Word 说明校验
    if files_provided.get("explanation", False):
        explanation_annotations = annotations.get("explanation", [])
        if not explanation_annotations:
            errors.append("上传了评估说明 Word 文件，但 annotations.explanation 为空。请提供说明章节的批注内容。")
        else:
            for i, ann in enumerate(explanation_annotations):
                original_text = ann.get("original_text", "")
                if not original_text:
                    errors.append(f"说明批注[{i}] 缺少 original_text 字段（应包含说明中的原文段落）")
                if not ann.get("description"):
                    errors.append(f"说明批注[{i}] 缺少 description 字段")
                if not ann.get("suggestion"):
                    errors.append(f"说明批注[{i}] 缺少 suggestion 字段")

    # 如果没有上传任何文件
    if not any(files_provided.values()):
        errors.append("请至少上传一个文件（Excel 评估报表 / Word 评估报告 / Word 评估说明）")

    # 校验：如果上传了文件，必须至少有一个批注
    total_annotations = (
        len(annotations.get("excel", [])) +
        len(annotations.get("report", [])) +
        len(annotations.get("explanation", []))
    )
    if any(files_provided.values()) and total_annotations == 0:
        errors.append("已上传文件但 annotations 为空。请提供至少一条批注内容。")

    if errors:
        error_msg = "\n".join(errors)
        logger.error(f"校验失败：{error_msg}")
        raise HTTPException(status_code=400, detail=error_msg)


def process_excel_annotations(excel_path: str, annotations: List[Dict]) -> Tuple[str, List[Dict]]:
    """
    处理 Excel 批注

    Returns:
        (输出文件路径，未匹配的批注列表)
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"评估报表_审核版_{timestamp}.xlsx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)

    unmatched = []

    with ExcelAnnotator(excel_path) as annotator:
        for issue in annotations:
            location = issue.get("location", "")
            if "!" not in location:
                unmatched.append({
                    "file_type": "excel",
                    "location": location,
                    "description": issue.get("description", ""),
                    "suggestion": issue.get("suggestion", ""),
                    "reason": "location 格式错误，应为 'Sheet 名!单元格' 格式"
                })
                continue

            sheet_name, cell = location.split("!", 1)

            # 验证单元格坐标是否有效（如 B15, C5 等）
            import re
            if not re.search(r'^[A-Z]+\d+$', cell.upper()):
                unmatched.append({
                    "file_type": "excel",
                    "location": location,
                    "description": issue.get("description", ""),
                    "suggestion": issue.get("suggestion", ""),
                    "reason": f"单元格坐标无效：{cell}"
                })
                continue

            comment_text = f"{issue['description']}\n\n建议：{issue.get('suggestion', '')}"

            success, error = annotator.add_comment(
                sheet_name=sheet_name,
                cell=cell,
                comment=comment_text,
                author="审核 AI"
            )

            if not success:
                unmatched.append({
                    "file_type": "excel",
                    "location": location,
                    "description": issue.get("description", ""),
                    "suggestion": issue.get("suggestion", ""),
                    "reason": error
                })
                continue

            if issue.get("severity") == "高":
                success, error = annotator.highlight_cell(sheet_name, cell, "FFFF00")
                if not success:
                    logger.warning(f"高亮单元格失败 {sheet_name}!{cell}: {error}")

        annotator.create_audit_sheet(annotations)
        annotator.save(output_path)

    return output_filename, unmatched


def process_word_annotations(word_path: str, annotations: List[Dict], file_type: str = "报告") -> Tuple[str, List[Dict]]:
    """
    处理 Word 批注（在原文位置添加，移除末尾汇总）

    Returns:
        (输出文件路径，匹配失败的警告列表)
    """
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"评估{file_type}_审核版_{timestamp}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)

    with WordAnnotator(word_path) as annotator:
        # 在原文位置添加批注
        warnings = annotator.annotate_document(annotations, file_type)

        # 调试：打印匹配结果
        logger.info(f"Word {file_type} 批注处理完成：{len(annotations)} 条批注，{len(warnings)} 条匹配失败")
        for w in warnings:
            original_text = w.get("original_text", "")[:50] + "..." if len(w.get("original_text", "")) > 50 else w.get("original_text", "")
            logger.warning(f"  - {original_text}: {w.get('reason', '')}")

        # 不再添加末尾汇总
        annotator.save(output_path)

    return output_filename, warnings


@router.post("/extract")
async def extract_excel(
    excel_file: Optional[UploadFile] = File(None, description="评估报表 Excel"),
    excel_url: Optional[str] = Form(None, description="评估报表 Excel URL"),
    mode: str = Form("audit", description="提取模式：audit=精简审核模式，full=完整模式")
):
    """
    提取 Excel 评估报表数据（支持文件上传或 URL 下载）

    - mode=audit：精简模式，只提取评估审核必需数据（体积小，推荐）
    - mode=full：完整模式，提取所有数据包括公式/备注（体积大）

    - 方式 1：直接上传 Excel 文件（excel_file）
    - 方式 2：提供 Excel 文件 URL（excel_url），服务器会下载后处理
    """
    try:
        if not excel_file and not excel_url:
            raise HTTPException(status_code=400, detail="请提供 excel_file（文件上传）或 excel_url（URL 下载）")

        excel_path = save_temp_file(file=excel_file, file_url=excel_url, suffix=".xlsx")
        logger.info(f"保存临时文件：{excel_path}")

        # 提取数据（手动管理生命周期以确保文件句柄释放）
        extractor = ExcelExtractor(excel_path)
        try:
            if mode == "full":
                excel_data = extractor.extract_all()
            else:
                excel_data = extractor.extract_for_audit()
        finally:
            extractor.close()
            logger.info(f"Excel 提取完成，关闭工作簿")

        # 删除临时文件
        try:
            os.unlink(excel_path)
            logger.info(f"临时文件已删除：{excel_path}")
        except Exception as e:
            logger.warning(f"删除临时文件失败 {excel_path}: {e}")

        return {
            "success": True,
            "excel_data": excel_data,
            "message": "Excel 提取完成，Word 请由 Dify 处理"
        }

    except httpx.HTTPError as e:
        raise HTTPException(status_code=400, detail=f"URL 下载失败：{str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"提取失败：{str(e)}")


@router.post("/extract/word")
async def extract_word_text(
    file_url: str = Form(..., description="Word 文件 URL（评估报告或评估说明）")
):
    """
    提取 Word 文档文本内容（支持 URL 下载）

    适用于：
    - 评估报告.docx
    - 评估说明.docx

    返回段落列表、完整文本和表格（Markdown 格式）
    """
    try:
        # 下载文件
        logger.info(f"开始下载 Word 文件：{file_url[:80]}...")
        async with httpx.AsyncClient(timeout=120) as client:
            resp = await client.get(file_url)
            resp.raise_for_status()
            file_content = resp.content

        logger.info(f"下载完成，大小：{len(file_content)} 字节")

        # 保存到临时文件（在同步线程池中执行，避免异步 IO 问题）
        import asyncio
        loop = asyncio.get_event_loop()

        unique_id = uuid.uuid4().hex[:12]
        temp_path = os.path.join(TEMP_DIR, f"word_{unique_id}_download.docx")

        # 确保目录存在
        os.makedirs(TEMP_DIR, exist_ok=True)

        logger.info(f"下载完成，大小：{len(file_content)} 字节")

        # 文件写入和检查
        def write_file(path, content):
            f = open(path, "wb")
            try:
                f.write(content)
                f.flush()
                os.fsync(f.fileno())
            finally:
                f.close()
            return os.path.getsize(path)

        file_size = await loop.run_in_executor(None, write_file, temp_path, file_content)
        await asyncio.sleep(0.5)

        if not os.path.exists(temp_path):
            raise IOError(f"文件写入后立即消失：{temp_path}")

        logger.info(f"文件写入成功：{temp_path} ({file_size} 字节)")

        # 文件可访问性检查
        try:
            with open(temp_path, "rb") as f:
                check_size = os.path.getsize(temp_path)
            logger.info(f"文件检查成功：存在且可读（{check_size} 字节）")
        except Exception as e:
            logger.error(f"文件检查失败：{e}")
            raise

        try:
            # 从文件读取 Word
            logger.info(f"开始读取 Word 文件：{temp_path}")
            doc = Document(temp_path)
            logger.info(f"Word 文件读取成功，段落数：{len(doc.paragraphs)}，表格数：{len(doc.tables)}")

            paragraphs = []
            tables_markdown = []

            # 提取段落（跳过表格内的段落，避免重复）
            logger.info("开始提取段落...")
            in_table = False
            for p in doc.paragraphs:
                # 检查段落是否在表格内
                if p._element.xpath('./ancestor::w:tbl'):
                    continue  # 跳过表格内的段落

                text = p.text.strip()
                if text:
                    paragraphs.append(text)

            logger.info(f"段落提取完成：{len(paragraphs)} 个段落")

            # 提取表格（转为 Markdown）
            logger.info("开始提取表格...")
            for i, table in enumerate(doc.tables):
                md_rows = []
                for row_idx, row in enumerate(table.rows):
                    cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
                    md_row = "| " + " | ".join(cells) + " |"
                    md_rows.append(md_row)

                    # 在第一行后添加表头分隔线
                    if row_idx == 0:
                        separator = "|" + "|".join(["---"] * len(cells)) + "|"
                        md_rows.append(separator)

                tables_markdown.append({
                    "table_index": i,
                    "markdown": "\n".join(md_rows),
                    "row_count": len(table.rows),
                    "col_count": len(table.columns)
                })

            logger.info(f"表格提取完成：{len(tables_markdown)} 个表格")

            # 构建完整内容（在段落中插入表格标记）
            content_parts = paragraphs.copy()
            for table_info in tables_markdown:
                content_parts.append(f"\n[表格 {table_info['table_index']+1} 开始]\n{table_info['markdown']}\n[表格结束]\n")

            result = {
                "success": True,
                "content": "\n\n".join(content_parts),
                "paragraphs": paragraphs,
                "tables": tables_markdown,
                "paragraph_count": len(paragraphs),
                "table_count": len(tables_markdown),
                "message": "Word 文本提取完成"
            }
            logger.info("Word 提取成功")

        except Exception as e:
            # 确保清理临时文件
            logger.error(f"Word 提取失败：{e}", exc_info=True)
            if os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                    logger.info(f"临时文件已清理：{temp_path}")
                except Exception as clean_err:
                    logger.warning(f"清理临时文件失败：{clean_err}")
            raise e

        # 清理临时文件（在 return 之前）
        if os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except:
                pass

        return result

    except httpx.HTTPError as e:
        raise HTTPException(status_code=400, detail=f"URL 下载失败：{str(e)}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Word 提取失败：{str(e)}")


@router.post("/annotate")
async def annotate_reports(
    excel_file: Optional[UploadFile] = File(None, description="评估报表 Excel"),
    excel_url: Optional[str] = Form(None, description="评估报表 Excel URL"),
    report_file: Optional[UploadFile] = File(None, description="评估报告 Word"),
    report_url: Optional[str] = Form(None, description="评估报告 Word URL"),
    explanation_file: Optional[UploadFile] = File(None, description="评估说明 Word"),
    explanation_url: Optional[str] = Form(None, description="评估说明 Word URL"),
    audit_result: str = Form(..., description="LLM 审核结果 JSON")
):
    """
    根据审核结果添加批注（支持文件上传或 URL 下载）

    audit_result 格式:
    {
      "audit_conclusion": "通过/有条件通过/不通过",
      "score": 0-100,
      "annotations": {
        "excel": [...],
        "report": [...],
        "explanation": [...]
      },
      "summary": {...}
    }
    """
    try:
        # 解析审核结果
        audit_data = json.loads(audit_result)
        logger.info(f"解析 audit_result: {json.dumps(audit_data, ensure_ascii=False)[:500]}...")

        # 校验文件和批注匹配（支持 URL）
        files_provided = {
            "excel": excel_file is not None or excel_url is not None,
            "report": report_file is not None or report_url is not None,
            "explanation": explanation_file is not None or explanation_url is not None
        }

        # 校验前打印日志
        logger.info(f"文件提供情况：excel={files_provided['excel']}, report={files_provided['report']}, explanation={files_provided['explanation']}")
        logger.info(f"Excel 批注数：{len(audit_data.get('annotations', {}).get('excel', []))}")
        logger.info(f"报告批注数：{len(audit_data.get('annotations', {}).get('report', []))}")
        logger.info(f"说明批注数：{len(audit_data.get('annotations', {}).get('explanation', []))}")

        validate_annotations(audit_data, files_provided)

        annotations = audit_data.get("annotations", {})
        result = {
            "success": True,
            "annotated_files": {},
            "summary": {},
            "match_warnings": []
        }

        # Excel 批注处理
        excel_unmatched = []
        if (excel_file or excel_url) and annotations.get("excel", []):
            logger.info(f"开始处理 Excel 批注，共 {len(annotations['excel'])} 条")
            excel_path = save_temp_file(file=excel_file, file_url=excel_url, suffix=".xlsx")
            logger.info(f"Excel 文件已保存到：{excel_path}")
            output_filename, unmatched = process_excel_annotations(excel_path, annotations["excel"])
            excel_unmatched = unmatched
            result["annotated_files"]["excel"] = f"/api/v1/report/download/{output_filename}"
            result["summary"]["excel_comments"] = len(annotations["excel"]) - len(unmatched)
            result["summary"]["excel_highlights"] = len([
                i for i in annotations["excel"] if i.get("severity") == "高"
            ])
            logger.info(f"Excel 批注处理完成，{len(unmatched)} 条未匹配")
            # 清理临时文件
            try:
                os.unlink(excel_path)
            except:
                pass

        # Word 报告批注处理
        report_warnings = []
        report_unmatched = []
        if (report_file or report_url) and annotations.get("report", []):
            logger.info(f"开始处理 Word 报告批注，共 {len(annotations['report'])} 条")
            report_path = save_temp_file(file=report_file, file_url=report_url, suffix=".docx")
            logger.info(f"报告文件已保存到：{report_path}")
            output_filename, warnings = process_word_annotations(report_path, annotations["report"], "报告")
            logger.info(f"Word 报告批注处理完成，{len(warnings)} 条未匹配")
            result["annotated_files"]["report"] = f"/api/v1/report/download/{output_filename}"
            result["summary"]["report_annotations"] = len(annotations["report"])
            report_warnings = warnings
            # 提取未匹配的批注信息
            report_unmatched = [{
                "file_type": "report",
                "original_text": w.get("original_text", ""),
                "description": w.get("description", ""),
                "suggestion": w.get("suggestion", ""),
                "reason": w.get("reason", "")
            } for w in warnings if w.get("original_text")]
            # 清理临时文件
            try:
                os.unlink(report_path)
            except:
                pass
        elif (report_file or report_url) and not annotations.get("report", []):
            logger.warning("上传了评估报告但没有批注，跳过不生成文件")
            result["skipped_files"] = result.get("skipped_files", [])
            result["skipped_files"].append({
                "file": "report",
                "reason": "annotations.report 为空"
            })

        # Word 说明批注处理
        explanation_warnings = []
        explanation_unmatched = []
        if (explanation_file or explanation_url) and annotations.get("explanation", []):
            logger.info(f"开始处理 Word 说明批注，共 {len(annotations['explanation'])} 条")
            explanation_path = save_temp_file(file=explanation_file, file_url=explanation_url, suffix=".docx")
            logger.info(f"说明文件已保存到：{explanation_path}")
            output_filename, warnings = process_word_annotations(explanation_path, annotations["explanation"], "说明")
            logger.info(f"Word 说明批注处理完成，{len(warnings)} 条未匹配")
            result["annotated_files"]["explanation"] = f"/api/v1/report/download/{output_filename}"
            result["summary"]["explanation_annotations"] = len(annotations["explanation"])
            explanation_warnings = warnings
            # 提取未匹配的批注信息
            explanation_unmatched = [{
                "file_type": "explanation",
                "original_text": w.get("original_text", ""),
                "description": w.get("description", ""),
                "suggestion": w.get("suggestion", ""),
                "reason": w.get("reason", "")
            } for w in warnings if w.get("original_text")]
            # 清理临时文件
            try:
                os.unlink(explanation_path)
            except:
                pass
        elif (explanation_file or explanation_url) and not annotations.get("explanation", []):
            logger.warning("上传了评估说明但没有批注，跳过不生成文件")
            result["skipped_files"] = result.get("skipped_files", [])
            result["skipped_files"].append({
                "file": "explanation",
                "reason": "annotations.explanation 为空"
            })

        # 添加匹配失败警告（统一收集所有未匹配的批注）
        all_warnings = report_warnings + explanation_warnings + excel_unmatched
        all_unmatched = report_unmatched + explanation_unmatched + excel_unmatched
        if all_warnings:
            result["match_warnings"] = all_warnings
            result["warning_count"] = len(all_warnings)
        # 单独返回未匹配的批注信息（包含 location/original_text, description, suggestion, reason）
        if all_unmatched:
            result["unmatched_annotations"] = all_unmatched
            result["unmatched_count"] = len(all_unmatched)

        # 添加成功标记
        result["success"] = True
        logger.info(f"批注处理完成，成功 {result['summary'].get('excel_comments', 0) + result['summary'].get('report_annotations', 0) + result['summary'].get('explanation_annotations', 0)} 条，失败 {len(all_unmatched)} 条")

        return result

    except json.JSONDecodeError as e:
        logger.error(f"JSON 解析错误：{e}")
        raise HTTPException(status_code=400, detail="audit_result 格式错误，必须是有效的 JSON 字符串")
    except httpx.HTTPError as e:
        logger.error(f"文件下载错误：{e}")
        raise HTTPException(status_code=400, detail=f"文件下载失败：{str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"批注处理异常：{e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"批注失败：{str(e)}")


@router.get("/download/{filename}")
async def download_file(filename: str):
    """下载带批注的文件"""
    from fastapi.responses import FileResponse

    file_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        file_path,
        media_type="application/octet-stream",
        filename=filename
    )


@router.get("/rules")
async def get_rules():
    """获取审核规则"""
    return {
        "excel_rules": {
            "required_sheets": ["资产明细表", "负债明细表", "净资产表", "评估结果汇总表"],
            "calculation_checks": ["汇总=明细之和", "公式无错误"],
            "comment_checks": ["重要科目有备注说明"]
        },
        "annotation_rules": {
            "excel_location_format": "Sheet 名!单元格，如'资产明细表!C3'",
            "word_location_format": "章节名称，如'评估方法'、'特别事项说明'",
            "required_fields": ["location", "description", "suggestion", "severity"]
        }
    }


@router.get("/health")
async def health_check():
    """健康检查"""
    return {
        "status": "ok",
        "service": "dify-bridge-report-audit",
        "features": ["Excel 提取", "Excel 单元格批注", "Word 原生批注", "文件 - 批注匹配校验"],
        "validation": {
            "excel_required": "上传 Excel 时必须提供 annotations.excel",
            "report_required": "上传报告时必须提供 annotations.report",
            "explanation_required": "上传说明时必须提供 annotations.explanation"
        }
    }
