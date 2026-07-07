# -*- coding: utf-8 -*-
"""
Word 评估报告批注写回
使用 python-docx>=1.2.0 原生 add_comment() 方法 - 兼容 WPS 和 Microsoft Office
"""
from docx import Document
from typing import List, Dict, Any, Tuple
import datetime
import logging
import os
import re

logger = logging.getLogger(__name__)


class WordAnnotator:
    """Word 原生批注写回 - 使用 python-docx>=1.2.0 add_comment() 方法"""

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self.match_warnings = []

    def add_comment_to_paragraph(self, para_index: int, comment_text: str,
                                  author: str = "审核 AI", initials: str = "SHR"):
        """
        在指定段落添加原生批注
        """
        if para_index < 0 or para_index >= len(self.doc.paragraphs):
            logger.warning(f"段落索引超出范围：{para_index}")
            return False

        para = self.doc.paragraphs[para_index]

        # 如果段落没有 runs，添加一个空白 run 作为批注锚点
        if not para.runs:
            para.add_run()

        # 使用文档对象的 add_comment 方法添加批注
        try:
            self.doc.add_comment(
                runs=para.runs,
                text=comment_text,
                author=author,
                initials=initials
            )
            logger.info(f"✅ 在段落{para_index}添加批注")
            return True
        except Exception as e:
            logger.error(f"添加批注失败：{e}")
            return False

    def annotate_document(self, annotations: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """批量处理批注（使用 original_text 原文摘抄匹配）"""
        warnings = []

        for i, ann in enumerate(annotations):
            original_text = ann.get("original_text", "")
            description = ann.get("description", "")
            suggestion = ann.get("suggestion", "")

            # 通过原文摘抄查找段落
            para_index, found, match_type = self.find_paragraph_by_text(original_text)

            # 找不到段落时记录到失败集合
            if not found or para_index < 0:
                warning = {
                    "annotation_index": i,
                    "file_type": "word",
                    "original_text": original_text[:100] + "..." if len(original_text) > 100 else original_text,
                    "description": description,
                    "suggestion": suggestion,
                    "reason": "在文档中未找到匹配的原文段落"
                }
                warnings.append(warning)
                self.match_warnings.append(warning)
                logger.warning(f"❌ 未找到匹配的原文段落：{original_text[:50]}...")
                continue

            # 找到段落，添加批注
            comment_text = f"{description}\n\n建议：{suggestion}"
            success = self.add_comment_to_paragraph(para_index, comment_text, "审核 AI", "AI")

            if success:
                logger.info(f"✅ 找到原文段落（索引{para_index}，匹配方式：{match_type}），添加批注")
            else:
                logger.warning(f"⚠️ 在段落{para_index}添加批注失败")

        return warnings

    def find_paragraph_by_text(self, target_text: str) -> Tuple[int, bool, str]:
        """
        根据原文摘抄查找段落

        匹配顺序：
        1. 严格匹配（原文直接匹配）
        2. 去除首尾特殊字符后匹配（只去除首尾，保留中间原文）
        3. 按省略号"..."分段匹配（取任意一段长度>=10 的片段）

        Args:
            target_text: 原文摘抄（目标文本）

        Returns:
            (段落索引，是否找到，匹配方式)
        """
        if not target_text:
            return -1, False, "空文本"

        logger.info(f"开始匹配目标文本：{target_text[:50]}...")

        # ========== 方式 1：严格匹配 ==========
        for i, para in enumerate(self.doc.paragraphs):
            if target_text.strip() in para.text:
                logger.info(f"在第{i}段找到（严格匹配）")
                return i, True, "严格匹配"

        # ========== 方式 2：去除首尾特殊字符 ==========
        target_cleaned = self._strip_special_chars(target_text)

        if target_cleaned != target_text:
            logger.info(f"去除首尾特殊字符后：{target_cleaned[:50]}...")
            for i, para in enumerate(self.doc.paragraphs):
                para_cleaned = self._strip_special_chars(para.text)
                if target_cleaned in para_cleaned:
                    logger.info(f"在第{i}段找到（去除首尾特殊字符）")
                    return i, True, "去除首尾特殊字符"

        # ========== 方式 3：按省略号分段匹配 ==========
        if "..." in target_text:
            fragments = target_text.split("...")
            logger.info(f"检测省略号：切分为{len(fragments)}段")

            for idx, fragment in enumerate(fragments):
                fragment = fragment.strip()
                if len(fragment) >= 10:  # 片段长度>=10 才匹配
                    logger.info(f"尝试片段{idx+1}：'{fragment[:30]}...'")
                    for i, para in enumerate(self.doc.paragraphs):
                        if fragment in para.text:
                            logger.info(f"在第{i}段找到（分段匹配，片段{idx+1}）")
                            return i, True, "分段匹配"

        logger.warning(f"未找到匹配的目标文本")
        return -1, False, "未匹配"

    def _strip_special_chars(self, text: str) -> str:
        """
        去除首尾特殊字符（标点、符号、空白等），保留中间核心内容

        只去除首尾，不动中间内容！

        Args:
            text: 原始文本

        Returns:
            去除首尾特殊字符后的文本
        """
        if not text:
            return text

        # 移除首尾空白
        text = text.strip()

        # 找到首个中文字符或字母数字的位置
        start = 0
        for i, char in enumerate(text):
            if '一' <= char <= '鿿' or char.isalnum():
                start = i
                break

        # 找到末个中文字符或字母数字的位置
        end = len(text)
        for i in range(len(text) - 1, -1, -1):
            if '一' <= text[i] <= '鿿' or text[i].isalnum():
                end = i + 1
                break

        return text[start:end]

    def _keep_only_chinese(self, text: str) -> str:
        """
        只保留中文字符和字母数字，移除标点符号

        Args:
            text: 原始文本

        Returns:
            只保留中文和字母数字后的文本
        """
        if not text:
            return text

        result = []
        for char in text:
            if '一' <= char <= '鿿' or char.isalnum():
                result.append(char)

        return ''.join(result)

    def get_match_warnings(self) -> List[Dict[str, Any]]:
        return self.match_warnings

    def save(self, output_path: str):
        """保存文档"""
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        self.doc.save(output_path)
        logger.info(f"保存文件到：{output_path}")

    def close(self):
        pass

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()
